# -*- coding: utf-8 -*-
"""生成《智能选座与导航系统 技术文档与接口说明》Word 文档，保存到 D:\\。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"D:\智能选座与导航系统-技术文档与接口说明.docx"

doc = Document()
n = doc.styles['Normal']; n.font.name = '微软雅黑'; n.font.size = Pt(10.5)
n.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def style_run(r, size=10.5, bold=False, mono=False, color=None, italic=False):
    r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
    r.font.name = 'Consolas' if mono else '微软雅黑'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas' if mono else '微软雅黑')
    if color: r.font.color.rgb = RGBColor(*color)


def shade(p, fill="F2F2F2"):
    pPr = p._p.get_or_add_pPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def h1(t):
    p = doc.add_heading(level=1); r = p.add_run(t); style_run(r, size=16, bold=True, color=(0x1F, 0x3B, 0x73))


def h2(t):
    p = doc.add_heading(level=2); r = p.add_run(t); style_run(r, size=13, bold=True, color=(0x2E, 0x5C, 0x9A))


def para(t, bold=False, bullet=False):
    p = doc.add_paragraph(style='List Bullet' if bullet else None)
    r = p.add_run(t); style_run(r, size=10.5, bold=bold)
    return p


def code(title, body):
    if title:
        p = doc.add_paragraph(); r = p.add_run(title); style_run(r, size=10, bold=True, color=(0x1F, 0x3B, 0x73))
    for ln in body.strip('\n').split('\n'):
        p = doc.add_paragraph(); r = p.add_run(ln if ln else ' ')
        style_run(r, size=8.5, mono=True)
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0); shade(p)
    doc.add_paragraph()


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Light Grid Accent 1'
    for i, hv in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''; r = c.paragraphs[0].add_run(hv); style_run(r, size=9, bold=True)
    for row in rows:
        cs = t.add_row().cells
        for i, val in enumerate(row):
            if i >= len(cs): break
            cs[i].text = ''
            for j, sub in enumerate(str(val).split('\n')):
                p = cs[i].paragraphs[0] if j == 0 else cs[i].add_paragraph()
                r = p.add_run(sub); style_run(r, size=8.5, mono=sub.strip().startswith(('python','POST','GET','PUT','#','[','curl')))
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


def callout(t, fill="FFF3E0"):
    p = doc.add_paragraph(); r = p.add_run(t); style_run(r, size=10.5, bold=True, color=(0xB0, 0x30, 0x00)); shade(p, fill)

# ============ 封面 ============
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('智能选座与导航一体化系统'); style_run(r, size=24, bold=True, color=(0x1F, 0x3B, 0x73))
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('技术文档与接口说明'); style_run(r, size=16, color=(0x40, 0x40, 0x40))
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('含“有人/无人占用判定”核心代码 · 项目结构与接口 · 代码审查结论'); style_run(r, size=10.5, italic=True, color=(0x70, 0x70, 0x70))
doc.add_paragraph()
table(['版本', '日期', '说明'],
      [['v1.0', '2026-09', '综合整理项目技术栈、占用判定逻辑、文件接口与代码审查结论；测试 70 项全部通过']],
      widths=[3, 4, 10])

# ============ 一、项目总览 ============
h1('一、项目总览与技术栈')
para('本项目是一套“物联网感知的公共空间智能选座与导航系统”：通过座位上的红外/PIR 传感器实时感知“有人/无人”，结合预约、锁定（m/n 动态防抢）、行为分析，为用户提供选座、预约、室内导航等功能。')
para('技术栈：')
code('', '''后端： Flask + Flask-SQLAlchemy + Flask-SocketIO + PyMySQL（可选 SQLite）
数据库： MySQL（seat_navigation）；可选 Redis（锁定/行为缓存）
前端： Vue3 + Element Plus（模板用 Jinja2 服务端渲染 + Vue 客户端）
室外地图： folium；室内地图： Fabric.js 平面图 + 路网
传感器/AI： OpenCV（图像形态学/特征匹配/拼接/直线检测）、轻量语义分割（规划中）
硬件端： ESP32（Arduino 框架）+ HC-SR501 PIR / 红外避障传感器，经 PlatformIO 编译''')
para('整体架构：ESP32（每座一个）按周期上报原始红外读数 → 主系统 /api/sensor/report → 状态机判定占用/释放 → 落库 + WebSocket 推送 → 前端实时刷新。')
para('设备端只做“持续扫描 + 上报”，占用判定/锁定/预约/导航全部由主系统完成，实现软硬件解耦。')

# ============ 二、占用判定 ============
h1('二、“有人/无人”占用判定核心代码')
para('占用判定统一由主系统完成：报告为“两束红外同时为 1”即有人（占用）；连续 2 次“非满”即无人（释放，预约时段内除外）。判定逻辑在真机上报与模拟器回调中共用同一套规则，保证一致性。')

h2('2.1 真实上报判定（app.py → sensor_report()）')
code('app.py：sensor_report() 核心判定段', '''
    # 两束同时遮挡 => 有人
    both = (ir_front == 1 and ir_back == 1)
    ...
    if both:
        if seat.status in ('free', 'error'):
            seat.status = 'occupied'
            seat.occupied_since = now
            seat.lock_available_since = now + timedelta(minutes=Config.LOCK_M_DEFAULT)
        seat.consecutive_empty = 0
    else:
        seat.consecutive_empty += 1
        if seat.consecutive_empty >= 2 and seat.status in ('occupied', 'error'):
            active_r = Reservation.query.filter(
                Reservation.seat_id == seat_id,
                Reservation.status == 'pending',
                Reservation.start_time <= now,
                Reservation.end_time > now,
            ).first()
            if active_r:
                seat.consecutive_empty = 0        # 预约时段内不因物理无人释放
            else:
                seat.status = 'free'
                seat.current_user_id = None
                seat.occupied_since = None
                seat.lock_available_since = None
    # 设备离线判定：距上次上报 >24h => 异常
    if previous_scan and hours_since > 24 and seat.status != 'error':
        seat.status = 'error'; seat.error_since = now
    # 异常自动恢复
    if previous_scan and hours_since <= 24 and seat.error_since:
        seat.error_since = None''')

h2('2.2 模拟器回调判定（app.py → sensor_callback()）')
para('模拟器回调与真实上报规则完全一致（“与 /api/sensor/report 保持一致”），并额外支持“异常座位收到有人上报自动恢复”。见 app.py 的 _build_sensor_simulator()。')

h2('2.3 设备端上报（DEMO/src/main.cpp → reportSeat()）')
code('DEMO/src/main.cpp：reportSeat()（读取 PIR → 上报原始读数）', '''
    int ir_front = readSensor(IR_SENSOR_A_PIN) ? 1 : 0;   // 读 HC-SR501 A
    int ir_back  = readSensor(IR_SENSOR_B_PIN) ? 1 : 0;   // 读 HC-SR501 B
    ...
    if (cfg_seat_id > 0)              doc["seat_id"] = cfg_seat_id;
    else if (cfg_seat_label.length()) doc["seat_label"] = cfg_seat_label;
    doc["ir_front"] = ir_front;
    doc["ir_back"]  = ir_back;
    doc["device_id"] = deviceId();     // 用 WiFi MAC
    int code = http.POST(payload);     // POST {server}/api/sensor/report''')
para('说明：PIR（HC-SR501）检测到人体移动时输出 HIGH → 固件 cfg_ir_active_high=true；`readSensor` 按该电平判断“检测到人”记为 1。主系统按 (1,1) 判定占用。')
callout('⚠️ PIR 特性：HC-SR501 识别的是“人体移动”而非“静止占座”，人久坐不动可能短暂掉回 0。建议把 HC-SR501 延时电位器调到最小、跳线选 H，可缓解。')

# ============ 三、项目结构与接口 ============
h1('三、项目文件结构与各文件/接口作用')
h2('3.1 后端（D:\\MAX_xiangmu）')
table(['文件', '作用'],
      [['app.py', '后端主入口。注册全部路由：用户/场所/座位/预约/锁定/导航/传感器上报/设备管理/抽象视图/系统配置。含占用状态机、WebSocket 推送、传感器模拟器装配。'],
       ['config.py', '系统配置：数据库、Redis、上传、锁定参数（m/n/t）、AI 权重、传感器间隔/离线阈值、`SENSOR_SIMULATOR_AUTOSTART`（接真机时设 False）。'],
       ['models/user.py', '用户模型（学生/管理员，角色 role，审核状态）。'],
       ['models/building.py', '建筑 Building / 楼层 Floor / 座位 Seat 模型。Seat 含 status(free/occupied/locked/error)、ir_front/ir_back、ir_enabled、last_scan_time、consecutive_empty 等传感器字段。'],
       ['models/reservation.py', '预约 Reservation / 锁定记录 LockRecord 模型。'],
       ['models/sensor_data.py', '传感器原始上报记录 SensorData（seat_id, ir_front, ir_back, timestamp）。'],
       ['models/sensor_device.py', '硬件设备 SensorDevice：device_id(WiFi MAC)、seat_id、ir_active_high、report_interval_ms、last_seen、is_new（设备注册/在线/新设备提示）。'],
       ['utils/locking.py', '锁定机制（m/n 动态防抢、行为记录 BehaviorTracker）。'],
       ['utils/navigation.py', '室内路网 / 寻路 PathFinder / 导航服务（到座位最近路径）。'],
       ['utils/mapget.py', '地图构建/图层（室外 folium、室内叠加）。'],
       ['utils/recommendation.py', 'AI 推荐（距离/热度/偏好/拥挤度加权）。'],
       ['utils/sensor_simulator.py', '传感器模拟器（随机翻转座位，开发/演示用）。'],
       ['utils/image_preprocessor.py', '平面图图像预处理。'],
       ['view/auto_mapping.py', '手机拍摄自动建图（视频/多图→平面图 JSON），复用 vitalframe/framecut 并补充直线检测/合并/JSON 输出。'],
       ['view/vitalframe.py', '视频关键帧提取。'],
       ['view/framecut.py', '特征匹配/单应性/全景拼接。'],
       ['setup_db.py', '初始化数据库（建表）。']])

h2('3.2 前端')
table(['文件', '作用'],
      [['templates/base.html', '全局模板（导航、Vue/axios/socket.io 加载）。'],
       ['templates/index.html / login.html / register.html / profile.html', '首页、登录、注册、个人中心。'],
       ['templates/seat_map.html', '用户端座位图（实时状态、选座、预约、导航，WebSocket 刷新）。'],
       ['templates/reservation.html', '预约页。'],
       ['templates/navigation.html', '室内导航页。'],
       ['templates/uploading.html', '平面图上传 + 自动建图入口（调用 /api/admin/mapping/tasks）。'],
       ['templates/admin/*.html', '管理后台：dashboard、buildings、floor_plan、hardware、settings、behavior、approvals、seats_qrcodes。'],
       ['static/js/app.js', '全局工具（api 请求封装、showToast、座位颜色等）。'],
       ['static/js/api-client.js', 'fetch 版后端 API 客户端（替代 mock-api）。'],
       ['static/js/admin/hardware.js', '硬件/传感器调试面板 Vue 逻辑（模拟器开关、设备管理、每座传感器、手动模拟上报、新设备提示）。'],
       ['static/js/admin/*.js', '各管理页 Vue 逻辑。'],
       ['static/js/seat_map.js', '用户端座位图逻辑。'],
       ['static/css/*.css', '样式（base、animations、utilities、pages/*）。']])

h2('3.3 硬件端（DEMO）')
table(['文件', '作用'],
      [['DEMO/src/main.cpp', 'ESP32 固件：设备配置门户（ESP32-Config 热点+网页表单）、WiFi MAC 作 device_id、开机注册 + 周期拉配置、读取 PIR 上报 ir_front/ir_back。'],
       ['DEMO/platformio.ini', 'PlatformIO 配置（espressif32 + esp32dev + arduino + ArduinoJson7；platform_packages 指向本地 GitHub 框架）。'],
       ['DEMO/.vscode/c_cpp_properties.json', 'C/C++ IntelliSense includePath（指向 arduino-esp32 框架/工具链头文件，消除编辑器红波浪线）。']])

h2('3.4 测试')
table(['文件', '作用'],
      [['tests/conftest.py', '测试 app/client 夹具（内存 SQLite，绝不触碰真实 MySQL）。'],
       ['tests/test_mapping.py', '自动建图接口测试。'],
       ['tests/test_hardware_report.py', 'ESP32 上报接口测试（seat_id/seat_label/floor_id/释放/歧义）。'],
       ['tests/test_sensor_device.py', '设备注册/配置下发/面板设备管理测试。'],
       ['tests/test_hardware_panel.py', '硬件/传感器调试面板页面与接口测试。'],
       ['tests/test_api.py / test_checkin.py / test_locking.py / test_frontend_integration.py', '基础 API、签到、锁定、前端集成测试。']])

# ============ 四、关键 API 接口 ============
h1('四、关键 API 接口一览')
table(['接口', '方法', '作用'],
      [['/api/sensor/report', 'POST', '设备/手动上报红外读数 ir_front/ir_back（支持 seat_id / seat_label + floor_id，带 device_id 刷新设备在线）；运行占用状态机。'],
       ['/api/sensor/device/register', 'POST', '设备开机注册/心跳（新设备自动登记 is_new=True）；返回配置。'],
       ['/api/sensor/device_config', 'GET', '设备拉取配置（seat_label/ir_active_high/interval），面板改配置免重烧。'],
       ['/api/admin/sensor/devices', 'GET', '管理端设备列表（在线/绑定/配置/is_new）。'],
       ['/api/admin/sensor/devices/<id>', 'PUT', '管理端改设备（绑定座位/传感器类型/间隔）；保存后清除 is_new。'],
       ['/api/admin/sensor/overview', 'GET', '硬件调试面板数据：全局参数 + 模拟器状态 + 每座位实时状态。'],
       ['/api/admin/simulator/start|stop|occupy', 'POST', '启动/停止模拟器、手动模拟占用/释放。'],
       ['/api/admin/config', 'GET/PUT', '系统配置（锁定参数/AI权重/传感器间隔/离线阈值）。'],
       ['/api/admin/mapping/tasks', 'POST', '自动建图（视频/多图 -> 平面图）。'],
       ['/api/seats/**、/api/floors/**、/api/buildings/**', 'GET/PUT/POST/DELETE', '场所/楼层/座位管理。'],
       ['/api/lock/**、/api/reservation/**、/api/behavior/**、/api/navigation/**', 'POST/GET', '锁定、预约、行为分析、导航。']])

# ============ 五、代码审查结论 ============
h1('五、代码审查结论（已修 & 现状）')
para('测试结果：**70 项全部通过**（pytest），无回归。以下为本轮审查发现并修复的问题：')
table(['问题', '位置', '修复'],
      [['共线合并只按角度分组，将不同位置平行墙并成一条（矩形+隔断只剩 1 条线）', 'view/auto_mapping.py merge_colinear_lines', '改为“角度 + 法向截距”聚类，只合并真正共线碎片。'],
       ['去重把垂直墙与水平墙在转角误判为重复', 'view/auto_mapping.py filter_wall_lines', '仅对同方向（同为水平或垂直）的墙去重。'],
       ['/api/sensor/report 只认 seat_id，真机按 seat_label 上报会失败', 'app.py sensor_report', '新增 seat_label + floor_id 支持与参数校验。'],
       ['按标签上报时 seat_id 为 None，写 SensorData 触发 NOT NULL', 'app.py sensor_report', '解析后统一 seat_id = seat.id。'],
       ['ESP32 配置保存未真正写入 Flash（缺 prefs.begin()）', 'DEMO/src/main.cpp handleSave', '保存前 open NVS，保存后 flush。'],
       ['配置错（5G 网络/服务器地址错）后设备死循环、无法重新配置', 'DEMO/src/main.cpp loop()', 'WiFi 连失败 / 注册失败若干次后自动清配置回配置页。'],
       ['服务器进程多个、只绑 127.0.0.1 导致设备连不上', '运行方式', '统一以 `python app.py 0.0.0.0 5800` 单实例启动，防火墙放行 5800。'],
       ['传感器模拟器随服务器自动开启，与真机数据互相覆盖', 'config.py / .env', '`.env` 设 SENSOR_SIMULATOR_AUTOSTART=False（接真机时）。'],
       ['ESP32 连不上 5G 网络', 'ESP32 硬件限制', 'ESP32 仅支持 2.4GHz，需用 2.4G 网络（如 CMCC-efeN，去 -5G）。']])

h2('5.1 遗留/注意（非错误）')
para('・ PIR 识别“移动”而非“静止”，久坐可能短暂掉 0 —— 属传感器固有特性，可调延时/灵敏度缓解。', bullet=True)
para('・ 自动建图拼接结果可能含家具/文字噪声线条 —— 属演示可接受，后续可加语义分割辅助。', bullet=True)
para('・ 设备端服务器地址必须填电脑局域网 IP（如 192.168.1.8:5800），不能是 127.0.0.1（对设备而言是它自己）。', bullet=True)

# ============ 六、数据库模型 ============
h1('六、核心数据模型')
table(['模型/表', '关键字段', '说明'],
      [['Building 建筑', 'id, name, region', '一栋楼'],
       ['Floor 楼层', 'id, building_id, floor_number, name', '楼内楼层'],
       ['Seat 座位', 'id, floor_id, seat_label, status(free/occupied/locked/error), ir_front, ir_back, ir_enabled, last_scan_time, consecutive_empty, occupied_since, lock_available_since, current_user_id', '核心座位（含传感器状态）'],
       ['SensorData', 'id, seat_id, ir_front, ir_back, timestamp', '传感器原始上报记录'],
       ['SensorDevice 设备', 'id, device_id(WiFi MAC), seat_id, ir_active_high, report_interval_ms, last_seen, is_new', 'ESP32 硬件设备（配置/在线/新设备）'],
       ['Reservation 预约', 'id, seat_id, user_id, start_time, end_time, status', '预约（pending/checked_in/completed/cancelled/no_show）'],
       ['LockRecord 锁定', 'id, seat_id, record_id, start, end, status', '锁定记录'],
       ['User 用户', 'id, student_id, name, role, is_approved', '用户/管理员']])

# ============ 结尾 ============
doc.add_paragraph()
callout('附录：硬件接真机操作见 docs/烧录与PIR接线操作指南.docx；自动建图见 docs/frontend-mapping-tasks.md；原理/协议见 docs/hardware-integration.md。', "E8F0FE")

doc.save(OUT)
print("已生成：", OUT)
