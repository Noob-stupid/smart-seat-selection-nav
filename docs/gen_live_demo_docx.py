# -*- coding: utf-8 -*-
"""把 docs/hardware-live-demo.md 转成 Word(.docx)，便于打印/答辩。"""
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD = r"D:\MAX_xiangmu\docs\hardware-live-demo.md"
OUT = r"D:\MAX_xiangmu\docs\硬件演示与自测操作指南.docx"

doc = Document()
normal = doc.styles['Normal']
normal.font.name = '微软雅黑'
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def style_run(run, size=10.5, bold=False, mono=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = 'Consolas' if mono else '微软雅黑'
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), 'Consolas' if mono else '微软雅黑')
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade(p, fill="F2F2F2"):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def add_inline(p, text, size=10.5, mono=False):
    """支持 **bold** 内联。"""
    for part in re.split(r'(\*\*.+?\*\*)', text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            style_run(r, size=size, bold=True, mono=mono or part.startswith('`'))
        else:
            # 行内代码 `xxx`
            for sub in re.split(r'(`[^`]+`)', part):
                if not sub:
                    continue
                if sub.startswith('`') and sub.endswith('`'):
                    r = p.add_run(sub[1:-1])
                    style_run(r, size=size - 0.5, mono=True)
                else:
                    r = p.add_run(sub)
                    style_run(r, size=size, mono=mono)


lines = open(MD, encoding='utf-8').read().split('\n')
i = 0
in_code = False
code_lines = []
while i < len(lines):
    ln = lines[i]
    stripped = ln.strip()
    # 代码块
    if stripped.startswith('```'):
        if not in_code:
            in_code = True
            code_lines = []
        else:
            in_code = False
            for cl in code_lines:
                p = doc.add_paragraph()
                r = p.add_run(cl if cl else ' ')
                style_run(r, size=9, mono=True)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                shade(p)
            doc.add_paragraph()
        i += 1
        continue
    if in_code:
        code_lines.append(ln)
        i += 1
        continue
    if not stripped:
        i += 1
        continue
    # 分隔线
    if stripped == '---':
        p = doc.add_paragraph()
        r = p.add_run('─' * 40)
        style_run(r, size=9, color=(0xB0, 0xB0, 0xB0))
        i += 1
        continue
    # 表格
    if stripped.startswith('|') and i + 1 < len(lines) and re.match(r'^\s*\|[\s:\-|]+\|\s*$', lines[i + 1]):
        header = [c.strip() for c in stripped.strip('|').split('|')]
        i += 2
        rows = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
            rows.append(cells)
            i += 1
        t = doc.add_table(rows=1, cols=len(header))
        t.style = 'Light Grid Accent 1'
        for idx, hval in enumerate(header):
            cell = t.rows[0].cells[idx]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(hval)
            style_run(r, size=9.5, bold=True)
        for row in rows:
            cells = t.add_row().cells
            for idx, val in enumerate(row):
                if idx >= len(cells):
                    break
                cells[idx].text = ''
                p = cells[idx].paragraphs[0]
                add_inline(p, val, size=9.5)
        doc.add_paragraph()
        continue
    # 标题
    if stripped.startswith('## '):
        p = doc.add_heading(level=2)
        r = p.add_run(stripped[3:])
        style_run(r, size=14, bold=True, color=(0x1F, 0x3B, 0x73))
        i += 1
        continue
    if stripped.startswith('# '):
        p = doc.add_heading(level=1)
        r = p.add_run(stripped[2:])
        style_run(r, size=17, bold=True, color=(0x1F, 0x3B, 0x73))
        i += 1
        continue
    if stripped.startswith('### '):
        p = doc.add_heading(level=3)
        r = p.add_run(stripped[4:])
        style_run(r, size=12, bold=True, color=(0x2E, 0x5C, 0x9A))
        i += 1
        continue
    # 引用
    if stripped.startswith('> '):
        p = doc.add_paragraph()
        add_inline(p, stripped[2:], size=10.5)
        shade(p, "FFF3E0")
        i += 1
        continue
    # 列表
    if stripped.startswith('- '):
        p = doc.add_paragraph(style='List Bullet')
        add_inline(p, stripped[2:])
        i += 1
        continue
    # 有序列表（1. 2. / 1、2、）
    m = re.match(r'^(\d+)[.、]\s*(.*)$', stripped)
    if m:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        r = p.add_run(m.group(1) + '. ')
        style_run(r, size=10.5, bold=True)
        add_inline(p, m.group(2))
        i += 1
        continue
    # 普通段落
    p = doc.add_paragraph()
    add_inline(p, stripped)
    i += 1

doc.save(OUT)
print("已生成：", OUT)
