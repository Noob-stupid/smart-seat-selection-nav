# -*- coding: utf-8 -*-
"""把 docs/烧录与PIR接线操作指南.md 转成 Word(.docx)。"""
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD = r"D:\MAX_xiangmu\docs\烧录与PIR接线操作指南.md"
OUT = r"D:\MAX_xiangmu\docs\烧录与PIR接线操作指南.docx"

doc = Document()
n = doc.styles['Normal']; n.font.name = '微软雅黑'; n.font.size = Pt(10.5)
n.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def style_run(r, size=10.5, bold=False, mono=False, color=None):
    r.font.size = Pt(size); r.font.bold = bold
    r.font.name = 'Consolas' if mono else '微软雅黑'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas' if mono else '微软雅黑')
    if color: r.font.color.rgb = RGBColor(*color)


def shade(p, fill="F2F2F2"):
    pPr = p._p.get_or_add_pPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def add_inline(p, t, size=10.5):
    for part in re.split(r'(\*\*.+?\*\*)', t):
        if not part: continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2]); style_run(r, size=size, bold=True)
        else:
            for sub in re.split(r'(`[^`]+`)', part):
                if not sub: continue
                if sub.startswith('`') and sub.endswith('`'):
                    r = p.add_run(sub[1:-1]); style_run(r, size=size - 0.5, mono=True)
                else:
                    r = p.add_run(sub); style_run(r, size=size)


lines = open(MD, encoding='utf-8').read().split('\n')
i = 0; incode = False; code = []
while i < len(lines):
    ln = lines[i]; st = ln.strip()
    if st.startswith('```'):
        if not incode: incode = True; code = []
        else:
            incode = False
            for cl in code:
                p = doc.add_paragraph(); r = p.add_run(cl if cl else ' ')
                style_run(r, size=9, mono=True); p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0); shade(p)
            doc.add_paragraph()
        i += 1; continue
    if incode: code.append(ln); i += 1; continue
    if not st: i += 1; continue
    if st == '---':
        p = doc.add_paragraph(); r = p.add_run('─' * 40); style_run(r, size=9, color=(0xB0, 0xB0, 0xB0)); i += 1; continue
    if st.startswith('|') and i + 1 < len(lines) and re.match(r'^\s*\|[\s:\-|]+\|\s*$', lines[i + 1]):
        header = [c.strip() for c in st.strip('|').split('|')]; i += 2; rows = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            rows.append([c.strip() for c in lines[i].strip().strip('|').split('|')]); i += 1
        t = doc.add_table(rows=1, cols=len(header)); t.style = 'Light Grid Accent 1'
        for idx, hv in enumerate(header):
            cell = t.rows[0].cells[idx]; cell.text = ''; r = cell.paragraphs[0].add_run(hv); style_run(r, size=9.5, bold=True)
        for row in rows:
            cells = t.add_row().cells
            for idx, val in enumerate(row):
                if idx >= len(cells): break
                cells[idx].text = ''; add_inline(cells[idx].paragraphs[0], val, 9.5)
        doc.add_paragraph(); i += 1; continue
    if st.startswith('## '):
        p = doc.add_heading(level=2); r = p.add_run(st[3:]); style_run(r, size=14, bold=True, color=(0x1F, 0x3B, 0x73)); i += 1; continue
    if st.startswith('# '):
        p = doc.add_heading(level=1); r = p.add_run(st[2:]); style_run(r, size=17, bold=True, color=(0x1F, 0x3B, 0x73)); i += 1; continue
    if st.startswith('> '):
        p = doc.add_paragraph(); add_inline(p, st[2:]); shade(p, "FFF3E0"); i += 1; continue
    if st.startswith('- '):
        p = doc.add_paragraph(style='List Bullet'); add_inline(p, st[2:]); i += 1; continue
    m = re.match(r'^(\d+)[.、]\s*(.*)$', st)
    if m:
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(18)
        r = p.add_run(m.group(1) + '. '); style_run(r, size=10.5, bold=True); add_inline(p, m.group(2)); i += 1; continue
    p = doc.add_paragraph(); add_inline(p, st); i += 1

doc.save(OUT)
print("已生成：", OUT)
